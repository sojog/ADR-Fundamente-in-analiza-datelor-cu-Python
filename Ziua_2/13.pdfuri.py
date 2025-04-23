# pip install python-docx pypdf2
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
import os

def get_pages_count(path):
    pdf = PdfReader(path)
    return len(pdf.pages)


def cut_pdf(path, start_page, stop_page, outputpath="outputfile.pdf"):
    pdf = PdfReader(path)
    pdf_writer = PdfWriter()
    for i in range(start_page - 1, stop_page):
        page = pdf.pages[i]
        pdf_writer.add_page(page)

    with open(outputpath, "wb") as f:
        pdf_writer.write(f)


def merge_pdfs(paths, outputpath="outputfile.pdf"):
    pdf_writer = PdfWriter()
    for path in paths:
        pdf = PdfReader(path)
        for page in pdf.pages:
            pdf_writer.add_page(page)
    with open(outputpath, "wb") as f:
        pdf_writer.write(f)


# export pdf to docx
def export_pdf_to_docx(path, outputpath="outputfile.docx"):
    """
    Convertește un fișier PDF în DOCX prin extragerea textului.
    Notă: Formatarea complexă, imaginile și alte elemente pot să nu fie 
    transferate corespunzător.
    """
    pdf = PdfReader(path)
    doc = Document()
    
    # Adăugăm titlu document
    doc.add_heading(f'Conținut convertit din {os.path.basename(path)}', 0)
    
    # Extragem textul din fiecare pagină și-l adăugăm în document
    for i, page in enumerate(pdf.pages):
        text = page.extract_text()
        
        # Adăugăm un heading pentru fiecare pagină
        doc.add_heading(f'Pagina {i+1}', level=1)
        
        # Adăugăm textul extras
        doc.add_paragraph(text)
        
        # Adăugăm un separator între pagini
        if i < len(pdf.pages) - 1:
            doc.add_paragraph('---')
    
    # Salvăm documentul
    doc.save(outputpath)
    print(f"Documentul a fost convertit și salvat ca {outputpath}")

# Testăm funcțiile
print(get_pages_count("Ziua_2/pdf_exemple.pdf"))
cut_pdf("Ziua_2/pdf_exemple.pdf", 2, 3, "pdf_exemple_2_3.pdf")
merge_pdfs(["pdf_exemple_2_3.pdf", "Ziua_2/pdf_exemple.pdf"], "pdf_exemple_2_3_4.pdf")

# Convertim PDF-ul în DOCX
export_pdf_to_docx("Ziua_2/pdf_exemple.pdf", "pdf_exemple.docx")










# if __name__ == "__main__":
#     path = "testfile.pdf"
#     cut_pdf(path, 4, 7, "hello.pdf")